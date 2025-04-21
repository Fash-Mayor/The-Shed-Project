from kivy.app import App
from kivy.uix.relativelayout import RelativeLayout
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.core.window import Window
from kivy.graphics import Rectangle, Color

from tkinter.filedialog import askopenfile
from PyPDF2 import PdfReader
from docx import Document

Window.size = (450, 600)

class SentenceLocatorApp(App):
    def fileChooser(self, event):
        #file selection dialogue box for choosing PDF and word files
        try:
            self.file = askopenfile(mode = "r", filetypes = [("supported files", "*.pdf *.docx"), ("pdf files", "*.pdf"), ("word files", "*.docx")])
            self.upload_file = self.file.name

            #display directory of uploaded file
            self.locationtext.text = self.upload_file
            self.locationtext.size_hint = (0.7, 0.11)
        except Exception as e:
            return f"Error reading file: {e}"
        
        #extract text from selected file
        self.text = ""
        #for pdf files
        if self.upload_file.endswith(".pdf"):
            try:
                with open(self.upload_file, 'rb') as file:
                    reader = PdfReader(file)
                    for page in reader.pages:
                        self.text += page.extract_text() + "\n"
            except FileNotFoundError:
                return "Error: PDF file not found."
            except Exception as e:
                return f"Error reading PDF: {e}"
            return self.text
        #for word files
        elif self.upload_file.endswith(".docx"):
            try:
                document = Document(self.upload_file)
                with open(self.upload_file,'rb') as file:
                    #reader = Document(file)
                    for paragraph in document.paragraphs:
                        self.text += paragraph.text + "\n"
            except FileNotFoundError:
                return "Error: Word file not found."
            except Exception as e:
                return f"Error reading Word file: {e}"
            return self.text
    
    def call_search_with_text(self, instance):
        search_text = self.findmeTextbox.text.lower()
        self.search(self.text, search_text)

    def search(self, text, search_text):
        # find_me_text = self.findmeTextbox.text.lower()
        self.searchingLabel.text = "==Searching=="

        if self.text is None:
            self.displayResultCount.text = "Occurences: 0"
            self.displayResultLocation.text = "Found on lines: No PDF loaded"

        lines = self.text.lower().splitlines()
        occurences = 0
        found_lines = []
        line_number = 1

        for line in lines:
            if search_text in line:
                occurences += line.count(search_text)
                found_lines.append(str(line_number))
                self.searchingLabel.text = "==Found=="
            line_number += 1 

        self.displayResultCount.text = f"Occurences: {occurences}"
        if found_lines:
            self.displayResultLocation.pos_hint = {"center_x": 0.31}
            self.displayResultLocation.text = f"Found on lines: {', '.join(found_lines)}"
        else:
            self.displayResultLocation.text = "Found on lines: No lines contain the search term."




    def build(self):
        layout = RelativeLayout()

        #changing backgroundcolor
        with layout.canvas.before:
            Color(0.8, 0.8, 0.8, 1)
            self.rect = Rectangle(pos = layout.pos, size = layout.size)

            layout.bind(pos = self._update_rect, size = self._update_rect)
        
        #widgets
        self.askopenfileButton = Button(text = "Choose File", size_hint = (None, None), height = 40,
                                        pos_hint = {"center_x": 0.14, "center_y": 0.9}, on_press = self.fileChooser)

        self.locationtext = TextInput(text ="", size_hint = (0.7, 0.07), pos_hint = {"center_x": 0.62, "center_y": 0.9},
                                      disabled = True)
        
        self.instructionLabel = Label(text = "Enter sentence or word to search for:", pos_hint = {"center_x": 0.4, "center_y": 0.8},
                                      font_name = "Tahoma", font_size = 20, bold = True, color = (0, 0, 0, 1))
        
        self.findmeTextbox = TextInput(size_hint = (0.9, 0.31), pos_hint = {"center_x": 0.48, "center_y": 0.6}, font_name = "Tahoma",
                                       font_size = 20)
        
        self.findmeButton = Button(text = "Search", size_hint = (None, None), height = 40, pos_hint = {"center_x": 0.14, "center_y": 0.39},
                                   on_press = self.call_search_with_text)
        
        self.searchingLabel = Label(text = "", size_hint = (1, 1), font_size = 24, font_name = "Tahoma", bold = True, color = (0, 0, 0, 1),
                                    pos_hint = {"center_x": 0.55, "center_y": 0.39})
        
        self.displayResultCount = Label(text = "Occurences: 0", size_hint = (None, None), height = 40, bold = True, color = (0, 0, 0, 1), 
                                        pos_hint = {"center_x": 0.2, "center_y": 0.26}, font_size = 25, font_name = "Tahoma")
        
        self.displayResultLocation = Label(text = "Found on lines:", size_hint = (None, None), height = 40, bold = True, color = (0, 0, 0, 1), 
                                           pos_hint = {"center_x": 0.217, "center_y": 0.15}, font_size = 25, font_name = "Tahoma")
        

        
        #adding widgets to layout
        layout.add_widget(self.displayResultLocation)
        layout.add_widget(self.displayResultCount)
        layout.add_widget(self.searchingLabel)
        layout.add_widget(self.findmeButton)
        layout.add_widget(self.findmeTextbox)
        layout.add_widget(self.instructionLabel)
        layout.add_widget(self.locationtext)
        layout.add_widget(self.askopenfileButton)
        return layout
    
    #changing backgroundcolor according to the resizing of screen. Kinda like responsiveness but for backgroundcolor 
    def _update_rect(self, instance, value):
        self.rect.pos = instance.pos
        self.rect.size = instance.size

if __name__ == "__main__":
    SentenceLocatorApp().run()